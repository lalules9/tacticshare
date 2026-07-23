"""
extract_confidence.py
─────────────────────
One-time script. Run this once to produce data/confidence_lookup.json.
After that, the app reads the JSON — you never need to run this again
unless you update list_b.docx.

Usage:
    python extract_confidence.py

Output:
    data/confidence_lookup.json
"""

import re
import json
import sys
from docx import Document

LIST_B_PATH = "data/list_b.docx"
OUTPUT_PATH = "data/confidence_lookup.json"

# The 14 structural variables in order.
# Each entry: (json_key, patterns_that_identify_this_variable_at_the_start_of_a_line)
VARIABLES = [
    ("target_type",                   ["target type"]),
    ("decision_making_concentration", ["concentration of decision", "decision-making power",
                                       "decision making power", "concentration of decision-making"]),
    ("strength_of_opposition",        ["strength of opposition"]),
    ("type_of_opposition",            ["type of opposition"]),
    ("visibility_of_harm",            ["visibility of harm", "visibility"]),
    ("behaviour_change_required",     ["degree of behaviour", "degree of behavior",
                                       "behaviour change required", "behavior change required"]),
    ("primary_mechanism",             ["role of policy", "policy vs market", "market pressure vs"]),
    ("coalition",                     ["coalition size", "coalition"]),
    ("economic_disruption",           ["economic disruption"]),
    ("speed_of_change",               ["speed of change", "speed achieved", "speed"]),
    ("scale_of_change",               ["scale of change", "scale achieved", "scale of impact", "scale"]),
    ("public_sentiment",              ["public sentiment", "public sentient"]),
    ("legal_regulatory",              ["legal/regulatory", "legal regulatory",
                                       "legal and regulatory", "legal landscape"]),
    ("narrative_dominance",           ["narrative dominance"]),
]

# Stop words that mean a line is NOT a variable heading even if it contains the keyword.
# e.g. "evidence" appearing in text about "speed" should not be treated as a heading.
NOT_A_HEADING = ["evidence:", "source tag:", "source:", "confidence:", "coding —",
                 "coding:", "note:", "hindsight", "variables coded"]


def extract_text(path):
    """
    Extract all text from the Word doc in document order, including table cells.
    python-docx's doc.paragraphs silently skips table content — this function
    walks the raw XML to get paragraphs and tables in the order they appear.
    """
    from docx.oxml.ns import qn
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    doc = Document(path)
    lines = []

    for child in doc.element.body.iterchildren():
        # Plain paragraph
        if child.tag == qn('w:p'):
            lines.append(Paragraph(child, doc).text)
        # Table — extract each cell's text
        elif child.tag == qn('w:tbl'):
            table = Table(child, doc)
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        if para.text.strip():
                            lines.append(para.text)

    return "\n".join(lines)


def split_into_cases(text):
    """Split full LIST B text into per-case sections keyed by case number (string)."""
    pattern = r'(?:LOCKED\s+)?CASE\s*#?(\d+)[:\s*–—*]+'
    parts = re.split(pattern, text, flags=re.IGNORECASE)
    cases = {}
    i = 1
    while i < len(parts) - 1:
        case_num = parts[i].strip()
        case_body = parts[i + 1]
        cases[case_num] = case_body
        i += 2
    return cases


def line_starts_with_variable(line):
    """
    Return the json_key if this line begins with a variable name, else None.
    Works regardless of line length — the variable name just needs to appear
    at the start after stripping bullet markers and asterisks.
    """
    # Strip leading bullets, asterisks, dashes, whitespace
    stripped = re.sub(r'^[\s\-\*•]+', '', line).strip()
    stripped_lower = stripped.lower()

    # Reject lines that are clearly evidence/source/confidence prose
    for bad in NOT_A_HEADING:
        if stripped_lower.startswith(bad):
            return None

    # Try to match each variable
    for json_key, patterns in VARIABLES:
        for pat in patterns:
            if stripped_lower.startswith(pat):
                return json_key

    return None


def find_confidence_in_text(text):
    """
    Find the first High/Medium/Low confidence statement in a block of text.
    Handles formats like:
      Confidence: High — reason
      Confidence: High. reason
      Confidence — High
    """
    for line in text.split('\n'):
        stripped = line.strip()
        if re.match(r'confidence', stripped, re.IGNORECASE):
            m = re.search(r'\b(High|Medium|Low)\b', stripped, re.IGNORECASE)
            if m:
                return m.group(1).capitalize()
    return None


def extract_variable_confidences(case_text):
    """
    Walk through the case text and extract confidence levels for each variable.
    Strategy: find lines that start with a variable name, then collect text
    until the next variable name or section break, then extract confidence.
    """
    lines = case_text.split('\n')

    # Build a list of (line_index, variable_key) for all detected variable headings
    variable_positions = []
    for i, line in enumerate(lines):
        key = line_starts_with_variable(line)
        if key:
            variable_positions.append((i, key))

    if not variable_positions:
        return {}

    # For each variable, collect text from its heading to the next heading (or end)
    result = {}
    for idx, (start_line, key) in enumerate(variable_positions):
        end_line = variable_positions[idx + 1][0] if idx + 1 < len(variable_positions) else len(lines)
        block = '\n'.join(lines[start_line:end_line])
        conf = find_confidence_in_text(block)
        if conf:
            result[key] = conf
        # If same key appears twice (e.g. "speed" matching both "speed" and "speed of change"),
        # keep only the first occurrence
        # (already handled since we don't overwrite)

    return result


def extract_overall_confidence(case_text):
    """
    Extract the overall case-level confidence flag.
    Looks for the OVERALL CASE CONFIDENCE section and the High/Medium/Low on that line
    or the next non-empty line.
    """
    lines = case_text.split('\n')
    for i, line in enumerate(lines):
        if re.search(r'OVERALL CASE CONFIDENCE', line, re.IGNORECASE):
            # Check this line first
            m = re.search(r'\b(High|Medium|Low)\b', line, re.IGNORECASE)
            if m:
                return m.group(1).capitalize()
            # Check next few lines
            for j in range(i + 1, min(i + 4, len(lines))):
                m = re.search(r'\b(High|Medium|Low)\b', lines[j], re.IGNORECASE)
                if m:
                    return m.group(1).capitalize()
    return "Not found"


def extract_campaign_name(case_text):
    """Extract campaign name from the first meaningful line of the case section."""
    for line in case_text.split('\n'):
        clean = re.sub(r'\*+', '', line).strip()
        clean = re.sub(r'\s*[\(–—]\s*\d{4}.*$', '', clean).strip()
        if clean and len(clean) > 5 and not re.match(r'^\d', clean):
            return clean
    return "Unknown"


def main():
    print(f"Reading {LIST_B_PATH}...")
    try:
        text = extract_text(LIST_B_PATH)
    except FileNotFoundError:
        print(f"ERROR: {LIST_B_PATH} not found.")
        sys.exit(1)

    print("Splitting into case sections...")
    cases = split_into_cases(text)
    print(f"Found {len(cases)} cases: {sorted(cases.keys(), key=int)}\n")

    results = {}
    all_ok = True

    for case_num in sorted(cases.keys(), key=int):
        body = cases[case_num]
        name = extract_campaign_name(body)
        overall = extract_overall_confidence(body)
        var_confidences = extract_variable_confidences(body)

        # Fill in any missing variables
        for key, _ in VARIABLES:
            if key not in var_confidences:
                var_confidences[key] = "Not found"

        results[case_num] = {
            "campaign_name": name,
            "variables": var_confidences,
            "overall_confidence": overall
        }

        missing = [k for k, v in var_confidences.items() if v == "Not found"]
        status_parts = [f"Overall: {overall}"]
        if missing:
            status_parts.append(f"⚠ Missing: {missing}")
            all_ok = False

        print(f"  Case {case_num}: {name[:55]}")
        for part in status_parts:
            print(f"    {part}")

    print(f"\nWriting {OUTPUT_PATH}...")
    with open(OUTPUT_PATH, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2, ensure_ascii=False)

    print(f"\nDone. {len(results)} cases written.")

    if all_ok:
        print("\n✓ No missing variables — extraction looks clean.")
    else:
        print("\n⚠ Some variables were not found. This means the parser couldn't detect")
        print("  the confidence statement for that variable in LIST B.")
        print("  Check the formatting of those sections in list_b.docx.")
        print("  The variable heading may not be on its own line, or 'Confidence:' may")
        print("  be spelled differently. Fix in the doc and re-run, or manually edit")
        print("  the JSON output for those entries.")


if __name__ == "__main__":
    main()